from fastapi import FastAPI, HTTPException, Body, Response, File, UploadFile, Query, Depends, status, Form, Request, Header
from fastapi.security import OAuth2PasswordBearer
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from fastapi.responses import JSONResponse
from nbformat import from_dict, write, read
from nbformat.v4 import new_code_cell, new_markdown_cell
from nbconvert.preprocessors import ExecutePreprocessor, CellExecutionError
#from nbconvert import PythonExporter, NotebookExporter

from traceback import format_exception

from typing import Dict, List, Optional, Union
import json
from dotenv import load_dotenv
import os
import shutil
import subprocess
import traceback
import logging
#import sys
import requests
import re

import boto3

import pickle

#import pdfplumber

load_dotenv()

from supabase import create_client, Client
#from starlette.responses import StreamingResponse
#import aiohttp
#import aiofiles
from fastapi import HTTPException
#from starlette.status import HTTP_400_BAD_REQUEST

from pydantic import BaseModel
from typing import Any 

import pandas as pd

import urllib

from urllib.parse import quote_plus

from collections import OrderedDict

from io import BytesIO
from docx import Document
from striprtf.striprtf import rtf_to_text

from langchain.text_splitter import RecursiveCharacterTextSplitter
from sentence_transformers import SentenceTransformer
from pydantic import BaseModel
from typing import Optional
#from urllib.parse import urlparse
#from urllib.parse import unquote
import PyPDF2

import uuid
from datetime import datetime

#from langchain.document_loaders import YoutubeLoader
from langchain_community.document_loaders import YoutubeLoader

import ast
from  RestrictedPython import compile_restricted

psqlpass = os.getenv("PSQLPASS")
#psqlpass = urllib.parse.quote_plus(psqlpass)
psqlpass = urllib.parse.quote_plus(str(psqlpass))

TOKEN = os.getenv("TOKEN")

API_KEY = os.getenv("API_KEY")

oauth2_scheme = OAuth2PasswordBearer(tokenUrl="token")

async def verify_token(token: str = Depends(oauth2_scheme)):
    if token != TOKEN:
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Invalid token",
            headers={"WWW-Authenticate": "Bearer"},
        )
    return token


async def verify_api_key(x_api_key: str = Header(...)):
    if x_api_key != API_KEY:
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Invalid API Key",
        )

class FolderName(BaseModel):
    folder_name: str

class ReplaceNotebookData(BaseModel):
    folder_name: str
    file_name: str
    content: Any

class DeleteFileInput(BaseModel):
    file_name: str

class ExecuteCellInput(BaseModel):
    folder_name: str
    file_name: str
    cell_index: int

class DownloadFileInput(BaseModel):
    url: str
    folder_name: str
    file_name: str

class ExecuteNotebook(BaseModel):
    folder_name: str
    file_name: str

class DLFileData(BaseModel):
    bucket_name: str
    folder_name: str
    file_name: str
    expiry_duration: int

class ULFileData(BaseModel):
    bucket_name: str
    folder_name: str
    file_name: str
    expiry_duration: int

class Cell(BaseModel):
    cell_type: Optional[str] = "code"
    execution_count: Optional[None] = None
    metadata: Optional[dict] = {}
    outputs: Optional[List] = []
    source: Optional[List[str]] = [""]

class UpdateCellInput(BaseModel):
    file_name: str
    cell_index: int
    cell_content: str
    cell_type: str = "code"

class DeleteCellInput(BaseModel):
    file_name: str
    cell_index: int

class Content(BaseModel):
    cells: List[Cell]
    metadata: Optional[dict] = {}
    nbformat: Optional[int] = 4
    nbformat_minor: Optional[int] = 4

class NotebookData(BaseModel):
    content: Content

class Notebook(BaseModel):
    file_name: Optional[str] = "default.ipynb"
    folder_name: Optional[str] = "default"
    notebook_data: NotebookData

class ExecuteNotebookInput(BaseModel):
    notebook_name: str

class FileName(BaseModel):
    file_name: str

class UploadFileInput(BaseModel):
    folder_path: str

class Chunks(BaseModel):
    chunk_size: int
    chunk_overlap: int

# class InputModel(BaseModel):
#     text: str
#     chunks: Chunks
#     return_full_text: bool
#     type: str
#     model: Optional[str] = None
class InputModel(BaseModel):
    text: Optional[str] = None
    url: Optional[str] = None
    type: str
    chunks: Chunks
    return_full_text: Optional[bool] = False
    model: Optional[str] = None

class InputModelURL(BaseModel):
    url: str
    chunks: Chunks
    return_full_text: bool = True
    type: str
    model: Optional[str] = None

class TranscribeInput(BaseModel):
    input_path: str
    output_path: str

url = os.getenv('SUPABASE_URL')
key = os.getenv('SUPABASE_KEY')
supabase: Client = create_client(url, key)

print(url)
print(key)

app = FastAPI()
origins = [
    "https://codebox.isolutionsai.com",
    "https://codebox.tcfcommandpost.org",
    # add more origins if needed
]


#environment set up
workspace_path = f"/home/ubuntu/automatenb/environment/"

app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)
def check_if_file_exists(file_path: str):
    if os.path.isfile(file_path):
        return True
    else:
        return False

def get_output_from_cell(nb, cell_index):
    cell = nb.cells[cell_index]
    if cell.cell_type == 'code':
        return cell.outputs
    else:
        return None

class FolderName(BaseModel):
    folder_name: str

def handle_input(notebook: Notebook):
    return notebook

@app.post("/list-files")
async def list_files(input: FolderName, token: str = Depends(verify_token)):
    try:
        folder_path = f"{workspace_path}/{input.folder_name}"
        files = os.listdir(folder_path)
        return {"status": "success", "files": files}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/read-notebook", response_class=Response)
async def read_notebook(file_name: FileName = Body(...), token: str = Depends(verify_token)):
    notebook_path = f"{workspace_path}/{file_name.file_name}"
    try:
        with open(notebook_path) as f:
            nb = read(f, as_version=4)

        # Convert the notebook to a dictionary
        nb_dict = dict(nb)

        # Extract only the cells part of the notebook
        cells = nb_dict['cells']

        # For each cell, keep only the 'outputs' and 'source' fields
        cells = [{'outputs': cell['outputs'], 'source': cell['source']} for cell in cells]

        # Convert the cells to a JSON string
        cells_json = json.dumps({"notebook": {"cells": cells}})

        # Return the JSON string as a stream
        #return Response(content=nb_json, media_type="application/json")
        return Response(content=cells_json, media_type="application/json")

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))



def execute_notebook(notebook_path, working_dir, cell_index):
    with open(notebook_path) as f:
        nb = read(f, as_version=4)

    # Add a code cell at the beginning of the notebook that changes the working directory
    change_dir_cell = new_code_cell(source=f'import os\nos.chdir(r"{working_dir}")')
    nb.cells.insert(0, change_dir_cell)

    ep = ExecutePreprocessor(timeout=600, kernel_name='python3')

    try:
        ep.preprocess(nb)
    except CellExecutionError as e:
        # Remove the change_dir_cell after execution
        nb.cells.pop(0)
        # Get the output from the cell at cell_index
        output = nb.cells[cell_index].outputs
        cells = []
        for cell in nb.cells:
            processed_outputs = []
            for output in cell.outputs:
                if 'traceback' in output:
                    output['traceback'] = [re.sub(r'\x1b\[.*?m', '', line) for line in output['traceback']]
                if output['output_type'] == 'display_data' and 'image/png' in output['data']:
                    del output['data']['image/png']
                processed_outputs.append(output)
            cells.append({
                "outputs": processed_outputs,
                "source": cell.source
            })
        tb_str = traceback.format_exc()
        # Remove escape characters from traceback
        tb_str = re.sub(r'\x1b\[.*?m', '', tb_str)
        return {"status": "error", "message": "Error executing cell", "traceback": tb_str, "notebook": {"cells": cells}, "output": output}

    # Remove the change_dir_cell after execution
    nb.cells.pop(0)

    # Get the output from the cell at cell_index
    output = nb.cells[cell_index].outputs

    # Prepare cells for output
    cells = []
    for cell in nb.cells:
        processed_outputs = []
        for output in cell.outputs:
            if output['output_type'] == 'display_data' and 'image/png' in output['data']:
                del output['data']['image/png']
            processed_outputs.append(output)
        cells.append({
            "outputs": processed_outputs,
            "source": cell.source
        })

    return {"notebook": {"cells": cells}, "output": output}


def safety_check(python_code: str) -> dict[str, object]:
    """Check if Python code is safe to execute.
    This function uses common patterns and RestrictedPython to check for unsafe patterns in the code.

    Args:
        python_code: Python code to check
    Returns:
        Dictionary with "safe" (bool) and "message" (str) keys
    """
    result = {"safe": True, "message": "The code is safe to execute."}

    # Crude check for problematic code (os, sys, subprocess, exec, eval, etc.)
    unsafe_modules = {"os", "sys", "subprocess", "builtins"}
    unsafe_functions = {
        "exec",
        "eval",
        "compile",
        "open",
        "input",
        "__import__",
        "getattr",
        "setattr",
        "delattr",
        "hasattr",
    }
    dangerous_builtins = {
        "globals",
        "locals",
        "vars",
        "dir",
        "eval",
        "exec",
        "compile",
    }
    # this a crude check first - no need to compile the code if it's obviously unsafe. Performance boost.
    try:
        tree = ast.parse(python_code)
    except SyntaxError as e:
        return {"safe": False, "message": f"Syntax error: {str(e)}"}

    for node in ast.walk(tree):
        if (
            isinstance(node, ast.Call)
            and isinstance(node.func, ast.Name)
            and node.func.id in dangerous_builtins
        ):
            return {
                "safe": False,
                "message": f"Use of dangerous built-in function: {node.func.id}",
            }
        # Check for unsafe imports
        if isinstance(node, ast.Import) or isinstance(node, ast.ImportFrom):
            module_name = node.module if isinstance(node, ast.ImportFrom) else None
            for alias in node.names:
                if module_name and module_name.split(".")[0] in unsafe_modules:
                    return {
                        "safe": False,
                        "message": f"Unsafe module import: {module_name}",
                    }
                if alias.name.split(".")[0] in unsafe_modules:
                    return {
                        "safe": False,
                        "message": f"Unsafe module import: {alias.name}",
                    }
        # Check for unsafe function calls
        elif isinstance(node, ast.Call):
            if isinstance(node.func, ast.Name) and node.func.id in unsafe_functions:
                return {
                    "safe": False,
                    "message": f"Unsafe function call: {node.func.id}",
                }
            elif (
                isinstance(node.func, ast.Attribute)
                and node.func.attr in unsafe_functions
            ):
                return {
                    "safe": False,
                    "message": f"Unsafe function call: {node.func.attr}",
                }

    try:
        # Compile the code using RestrictedPython with a filename indicating its dynamic nature
        compiled_code = compile_restricted(
            python_code, filename="<dynamic>", mode="exec"
        )
        # Note: Execution step is omitted to only check the code without running it
        # This is not perfect, but should catch most unsafe patterns
    except Exception as e:
        return {
            "safe": False,
            "message": f"RestrictedPython detected an unsafe pattern: {str(e)}",
        }

    return result

@app.post("/execute-notebook")
async def execute(notebook_name: str = Body(...), safety_check_flag: bool = Body(False), api_key: str = Depends(verify_api_key)):
    try:
        folder_path = f"{workspace_path}"
        notebook_path = f"{folder_path}/{notebook_name}"
        cell_index = 0
        # Load the notebook and perform a safety check on each cell if the flag is set
        if safety_check_flag:
            with open(notebook_path) as f:
                nb = read(f, as_version=4)
            for cell in nb.cells:
                if cell.cell_type == 'code':
                    check_result = safety_check(cell.source)
                    if not check_result['safe']:
                        return {"status": "error", "message": check_result['message']}
        result = execute_notebook(notebook_path, folder_path, cell_index)
        if "status" in result and result["status"] == "error":
            return result
        return {"notebook": result["notebook"], "output": result["output"]}
    except Exception as e:
        logging.error(f"Error executing cell: {e}")
        raise HTTPException(status_code=500, detail=str(e))

### Return as a postgresql
from sqlalchemy import create_engine



@app.get("/read-pgsql")
async def read_pgsql(table_name: str = Query(...), token: str = Depends(oauth2_scheme), api_key: str = Depends(verify_api_key)):
    try:
        encoded_psqlpass = quote_plus(psqlpass)
        engine = create_engine(f'postgresql://{id}:{encoded_psqlpass}@aws-0-us-west-1.pooler.supabase.com:5432/postgres')
        query = f"SELECT * FROM {table_name} LIMIT 25"
        df = pd.read_sql_query(query, engine)
        return {"status": "success", "data": df.to_dict(orient='records')}
    except Exception as e:
        logging.error(f"Error reading from table: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/get-table-names")
async def get_table_names(token: str = Depends(oauth2_scheme)):
    try:
        encoded_psqlpass = quote_plus(psqlpass)
        engine = create_engine(f'postgresql://{id}:{encoded_psqlpass}@aws-0-us-west-1.pooler.supabase.com:5432/postgres')
        query = "SELECT table_name FROM information_schema.tables WHERE table_schema = 'public'"
        df = pd.read_sql_query(query, engine)
        return {"status": "success", "tables": df['table_name'].tolist()}
    except Exception as e:
        logging.error(f"Error getting table names: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/update-pkl")
async def read_pgsql(table_name: str = Query(...),psqlpass: str = Body(...),file_name: str = Body(...), sql: str = Body(...), token: str = Depends(oauth2_scheme), api_key: str = Depends(verify_api_key)):
    try:
        encoded_psqlpass = quote_plus(psqlpass)
        engine = create_engine(f'postgresql://{id}:{encoded_psqlpass}@aws-0-us-west-1.pooler.supabase.com:5432/postgres')


        df = pd.read_sql_query(sql, engine)

        df.to_pickle(f'{workspace_path}/{file_name}')

        return FileResponse(f'{file_name}', media_type='application/octet-stream')
    except Exception as e:
        logging.error(f"Error reading from table: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/write-pkl-supabase")
async def write_pkl_supabase(id: str = Body(...), psqlpass: str = Body(...),file_name: str = Body(...), sql: str = Body(...), api_key: str = Depends(verify_api_key)):
    try:
        encoded_psqlpass = quote_plus(psqlpass)
        engine = create_engine(f'postgresql://{id}:{encoded_psqlpass}@aws-0-us-west-1.pooler.supabase.com:5432/postgres')

        df = pd.read_sql_query(sql, engine)

        # Create the directory if it does not exist
        directory = os.path.dirname(f'{workspace_path}/{file_name}')
        if not os.path.exists(directory):
            os.makedirs(directory)

        df.to_pickle(f'{workspace_path}/{file_name}')

        return {"status": 200, "message": "file created", "filename": file_name}
    except Exception as e:
        logging.error(f"Error writing to pickle file: {e}")
        raise HTTPException(status_code=500, detail=str(e))



import numpy as np
from sqlalchemy import inspect
import pytz

def map_sqlalchemy_types_to_python(column_types):
    """
    Maps SQLAlchemy type names to Python types suitable for pandas DataFrames.
    
    Args:
    column_types (dict): A dictionary mapping column names to SQLAlchemy type names.
    
    Returns:
    dict: A dictionary mapping column names to Python types.
    """
    type_mapping = {
        "Integer": np.int64,
        "INTEGER": np.int32,  # Equivalent to int4 in PostgreSQL
        "BIGINT": np.int64,   # Equivalent to int8 in PostgreSQL
        "String": str,
        "Text": str,          # Text type in PostgreSQL
        "Float": np.float64,
        "Numeric": float,     # Numeric can be safely mapped to float for most use cases
        "Date": 'datetime64[ns]',
        "DateTime": 'datetime64[ns]',
        "Boolean": bool,
        "UUID": str,          # UUIDs are best represented as strings
        "TIMESTAMP": 'datetime64[ns]',
        "JSONB": str,          # JSONB can be represented as strings or objects
        "BOOLEAN": bool,
        "TEXT": str,  
        # Add more mappings as needed
    }
    
    python_types = {}
    for column, sql_type in column_types.items():
        # Get the name of the SQL type
        sql_type_name = sql_type.__class__.__name__
        # Find the closest Python type, default to str if not found
        python_type = type_mapping.get(sql_type_name, str)
        python_types[column] = python_type
    
    return python_types

@app.post("/write-pkl-supabase-datatypes")
async def write_pkl_supabase(id: str = Body(...), psqlpass: str = Body(...), file_name: str = Body(...), sql: str = Body(...), api_key: str = Depends(verify_api_key)):
    try:
        encoded_psqlpass = quote_plus(psqlpass)
        engine = create_engine(f'postgresql://{id}:{encoded_psqlpass}@aws-0-us-west-1.pooler.supabase.com:5432/postgres')

        # Execute SQL query to get the data
        df = pd.read_sql_query(sql, engine)
        
        # Get the table name from the SQL query for schema inspection
        table_name = sql.split("FROM")[1].strip().split()[0]  # Simplistic way to extract table name, might need improvement
        
        # Inspect the schema of the table
        inspector = inspect(engine)
        columns_info = inspector.get_columns(table_name)
        
        # Log the column information
        logging.info(f"Columns info: {columns_info}")
        
        # Create a dictionary to map column names to their SQL data types
        column_types = {col['name']: col['type'] for col in columns_info}
        
        # Log the column types
        logging.info(f"Column types: {column_types}")
        
        # Map SQLAlchemy types to Python types
        python_types = map_sqlalchemy_types_to_python(column_types)
        
        # Log the Python types
        logging.info(f"Python types: {python_types}")
        
        # Convert DataFrame columns to the corresponding Python data types
        for column in df.columns:
            if column in python_types:
                if python_types[column] == 'datetime64[ns]':
                    # Handle timezone-naive timestamps
                    if df[column].dt.tz is None:
                        df[column] = df[column].dt.tz_localize('UTC')
                    else:
                        # Handle timezone-aware timestamps
                        df[column] = df[column].dt.tz_convert('UTC')
                elif python_types[column] == bool:
                    # Handle boolean types
                    df[column] = df[column].astype('bool')
                elif python_types[column] == 'object' and column_types[column].__class__.__name__ == 'JSONB':
                    # Handle JSONB types
                    df[column] = df[column].apply(json.loads)
                else:
                    df[column] = df[column].astype(python_types[column])
                # Log the conversion
                logging.info(f"Converted column {column} to {python_types[column]}")

        # Convert any remaining 'object' type columns to 'str'
        for column in df.columns:
            if df[column].dtype == 'object':
                df[column] = df[column].astype(str)
                logging.info(f"Converted column {column} to str")

        # Save DataFrame and data types to a pickle file
        data = {
            'dataframe': df,
            'dtypes': {col: 'str' if df[col].dtype == 'object' else dtype for col, dtype in df.dtypes.apply(lambda x: x.name).to_dict().items()}
        }

        # Log the data types
        logging.info(f"Data types: {data['dtypes']}")

        # Create the directory if it does not exist
        directory = os.path.dirname(f'{workspace_path}/{file_name}')
        if not os.path.exists(directory):
            os.makedirs(directory)

        with open(f'{workspace_path}/{file_name}', 'wb') as f:
            pickle.dump(data, f)

        return {"status": 200, "message": "file created", "filename": file_name}
    except Exception as e:
        logging.error(f"Error writing to pickle file: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/update-pkl-from-file")
async def update_pkl_from_file(file: UploadFile = File(...), file_name: str = Body(...), api_key: str = Depends(verify_api_key)):
    try:
        # Determine the file type
        file_type = file.filename.split('.')[-1]

        if file_type not in ['csv', 'json']:
            raise HTTPException(status_code=400, detail="Invalid file type. Only CSV and JSON files are supported.")

        # Read the file into a pandas DataFrame
        if file_type == 'csv':
            df = pd.read_csv(file.file)
        else:  # json
            df = pd.read_json(file.file)
        # Create the directory if it does not exist
        directory = os.path.dirname(f'{workspace_path}/{file_name}')
        if not os.path.exists(directory):
            os.makedirs(directory)


        # Write the DataFrame to a pickle file
        df.to_pickle(f'{workspace_path}/{file_name}')

        return {"status": 200, "message": "Pickle file updated", "file_name": file_name}
    except Exception as e:
        logging.error(f"Error updating pickle file: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# @app.get("/read-pkl")
# async def read_pgsql(file_name: str = Query(...), token: str = Depends(oauth2_scheme)):
#     try:
#         # Load the DataFrame from the pickle file
#         df = pd.read_pickle(f'{workspace_path}/{file_name}')

#         # Convert the DataFrame to a dictionary
#         data = df.to_dict(orient='records')

#         return {"status": "success", "data": data}
#     except Exception as e:
#         logging.error(f"Error reading from pickle file: {e}")
#         raise HTTPException(status_code=500, detail=str(e))

@app.get("/read-pkl")
async def read_pgsql(file_name: str = Query(...), token: str = Depends(oauth2_scheme)):
    try:
        # Load the DataFrame and data types from the pickle file
        with open(f'{workspace_path}/{file_name}', 'rb') as f:
            data = pickle.load(f)
        
        # Check if 'dtypes' exists in the loaded data
        if 'dtypes' in data:
            df = data['dataframe']
            dtypes = data['dtypes']
            
            # Apply the data types to the DataFrame
            for column, dtype in dtypes.items():
                df[column] = df[column].astype(dtype)
        else:
            # If 'dtypes' does not exist, assume 'data' is the DataFrame
            df = data

        # Convert the DataFrame to a dictionary
        data = df.to_dict(orient='records')

        return {"status": "success", "data": data}
    except Exception as e:
        logging.error(f"Error reading from pickle file: {e}")
        raise HTTPException(status_code=500, detail=str(e))
### Not ready for use yet.  Write to supabase is out of scope currently.

@app.get("/modify-pgsql-pkl")
async def modify_pgsql(token: str = Depends(oauth2_scheme), api_key: str = Depends(verify_api_key)):
    try:
        # Load the DataFrame from the pickle file
        df = pd.read_pickle('cache.pkl')

        # Check if DataFrame is not empty
        if df.empty:
            return {"status": "error", "message": "Table is empty"}
        else:
            # Return the first 10 rows
            data = df.head(10).to_dict(orient='records')
            return {"status": "success", "message": "Table is not empty", "data": data}
    except Exception as e:
        logging.error(f"Error modifying pickle file: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/execute-notebook-pgsql")
async def execute(input: ExecuteNotebook, id: str = Body(...), psqlpass: str = Body(...), api_key: str = Depends(verify_api_key)):
    try:
        folder_path = f"{workspace_path}/{input.folder_name}"
        notebook_path = f"{folder_path}/{input.file_name}"
        cell_index = 0
        result = execute_notebook(notebook_path, folder_path, cell_index)
        if "status" in result and result["status"] == "error":
            return result

        encoded_psqlpass = quote_plus(psqlpass)
        engine = create_engine(f'postgresql://{id}:{encoded_psqlpass}@aws-0-us-west-1.pooler.supabase.com:5432/postgres')

        # Convert the output to a dataframe
        df = pd.DataFrame(result["output"])

        # Write the dataframe to a table in the database
        df.to_sql('list', engine, if_exists='replace')

        return {"status": "success", "message": "Output written to PostgreSQL table 'my_table'"}
    except Exception as e:
        logging.error(f"Error executing cell: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/update-supabase")
async def update_supabase(table_name: str = Body(...), token: str = Depends(oauth2_scheme)):
    try:
        # Load the DataFrame from the pickle file
        df = pd.read_pickle('cache.pkl')

        encoded_psqlpass = quote_plus(psqlpass)
        engine = create_engine(f'postgresql://{id}:{encoded_psqlpass}@aws-0-us-west-1.pooler.supabase.com:5432/postgres')

        # Write the DataFrame to a table in the database
        df.to_sql(table_name, engine, if_exists='replace')

        return {"status": "success", "message": f"Table '{table_name}' updated in Supabase"}
    except Exception as e:
        logging.error(f"Error updating Supabase table: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/new-notebook")
async def notebook_environment(notebook: dict = Body(...), file_name: str = Body(...), safety_check_flag: bool = Body(False), api_key: str = Depends(verify_api_key)):
    # Transform the request body
    cells = [{"cell_type": "code", "source": cell.split('\n')} for cell in notebook['cells']]
    notebook_data = {
        "execute": notebook['execute'],
        "file_name": file_name,
        "notebook_data": {
            "content": {
                "cells": cells
            }
        }
    }
    # Convert the transformed request body into a Notebook instance
    notebook = Notebook(**notebook_data)
    folder_name = f"{workspace_path}"
    notebook_content = notebook.notebook_data.content.dict()  # Convert to dictionary

    # Extract the directory from the file_name
    directory = os.path.join(workspace_path, os.path.dirname(file_name))

    # Create the directory if it does not exist
    if not os.path.exists(directory):
        os.makedirs(directory)

    notebook_path = os.path.join(folder_name, file_name)

    try:
        # Convert the list of strings into a single string for each cell
        for cell in notebook_content['cells']:
            cell['source'] = '\n'.join(cell['source'])
        # Convert the dictionary to a notebook object
        nb = from_dict(notebook_content)

        # Write the notebook node to a file
        with open(notebook_path, 'w') as f:
            write(nb, f)

        # Perform a safety check on each cell
        for cell in nb.cells:
            if cell.cell_type == 'code'and safety_check_flag:
                check_result = safety_check(cell.source)
                if not check_result['safe']:
                    return {"status": "error", "message": check_result['message']}
                
        # Execute the notebook
        result = execute_notebook(notebook_path, folder_name, 0)

        if "status" in result and result["status"] == "error":
            # Ensure that "output" is an array
            output = result["output"]
            if not isinstance(output, list):
                output = [output]
            return {"status": "error", "notebook": result["notebook"], "output": output}

        # Ensure that "output" is an array
        output = result["output"]
        if not isinstance(output, list):
            output = [output]

        return {"notebook": result["notebook"], "output": output}      
    except Exception as e:
        return {"status": "error", "output": {"cells": cells}, "detail": str(e).replace("\x1b", " ")}


@app.post("/delete-file")
async def delete_file(input: DeleteFileInput, api_key: str = Depends(verify_api_key)):
    file_path = f"{workspace_path}/{input.file_name}"
    print(file_path)
    if os.path.isfile(file_path):
        try:
            os.remove(file_path)
            return {"status": "success", "message": f"File {input.file_name} has been deleted."}
        except Exception as e:
            raise HTTPException(status_code=500, detail=str(e))
    else:
        return {"status": "error", "message": "File does not exist"}

@app.post("/upload-file")
async def upload_file(file_name: UploadFile = File(...), folder_path: str = Form(...), api_key: str = Depends(verify_api_key)):
    full_folder_path = f"{workspace_path}/{folder_path}"
    if not os.path.isdir(full_folder_path):
        os.makedirs(full_folder_path, exist_ok=True)

    with open(f"{full_folder_path}/{file_name.filename}", "wb") as buffer:
        shutil.copyfileobj(file_name.file, buffer)

    return {"file_name": file_name.filename}

#Uploads file to fs from any URL
@app.post("/upload-url")
async def download_file(input: DownloadFileInput, api_key: str = Depends(verify_api_key)):
    folder_name = f"{workspace_path}/{input.folder_name}"
    if not os.path.isdir(folder_name):
        os.makedirs(folder_name, exist_ok=True)
    file_path = f"{folder_name}/{input.file_name}"

    try:
        with requests.get(input.url, stream=True) as r:
            r.raise_for_status()
            with open(file_path, 'wb') as f:
                for chunk in r.iter_content(chunk_size=8192):
                    f.write(chunk)
        return {"filename": input.file_name}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/delete-folder")
async def delete_folder(folder_name: FolderName, api_key: str = Depends(verify_api_key)):
    folder_path = f"/home/ubuntu/automatenb/environment/{folder_name.folder_name}"
    if os.path.isdir(folder_path):
        try:
            shutil.rmtree(folder_path)
            return {"status": "success", "message": f"Folder {folder_name.folder_name} has been deleted."}
        except Exception as e:
            raise HTTPException(status_code=500, detail=str(e))
    else:
        return {"status": "error", "message": "Folder does not exist"}

@app.post("/add-cell")
async def add_cell(file_name: str = Body(...), cell_content: str = Body(...), cell_type: str = Body("code"), safety_check_flag: bool = Body(False), api_key: str = Depends(verify_api_key)):
    try:
        # Check if the code is safe to execute
        if safety_check_flag:
            check_result = safety_check(cell_content)
            if not check_result['safe']:
                return {"status": "error", "message": check_result['message']}


        notebook_path = f"{workspace_path}/{file_name}"
        with open(notebook_path) as f:
            nb = read(f, as_version=4)

        if cell_type == 'code':
            new_cell = new_code_cell(source=cell_content)
        elif cell_type == 'markdown':
            new_cell = new_markdown_cell(source=cell_content)
        else:
            return {"status": "error", "message": "Invalid cell type"}

        new_cell_index = len(nb.cells)
        nb.cells.append(new_cell)

        with open(notebook_path, 'w') as f:
            write(nb, f)

        return {"status": "success", "cell_index": new_cell_index}
    except Exception as e:
        logging.error(f"Error updating cell: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/update-cell")
async def update_cell(input: UpdateCellInput, safety_check_flag: bool = Body(False)):
    try:
        # Check if the code is safe to execute
        if safety_check_flag:
            check_result = safety_check(input.cell_content)
            if not check_result['safe']:
                return {"status": "error", "message": check_result['message']}

        notebook_path = f"{workspace_path}/{input.file_name}"
        with open(notebook_path) as f:
            nb = read(f, as_version=4)

        if input.cell_type == 'code':
            new_cell = new_code_cell(source=input.cell_content)
        elif input.cell_type == 'markdown':
            new_cell = new_markdown_cell(source=input.cell_content)
        else:
            return {"status": "error", "message": "Invalid cell type"}

        if input.cell_index < 0 or input.cell_index >= len(nb.cells):
            return {"status": "error", "message": "Invalid cell index"}

        nb.cells[input.cell_index] = new_cell

        with open(notebook_path, 'w') as f:
            write(nb, f)

        return {"status": "success"}
    except Exception as e:
        logging.error(f"Error updating cell: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/delete-cell")
async def delete_cell(input: DeleteCellInput, api_key: str = Depends(verify_api_key)):
    try:
        notebook_path = f"{workspace_path}/{input.file_name}"
        with open(notebook_path) as f:
            nb = read(f, as_version=4)

        if input.cell_index < 0 or input.cell_index >= len(nb.cells):
            return {"status": "error", "message": "Invalid cell index"}

        del nb.cells[input.cell_index]

        with open(notebook_path, 'w') as f:
            write(nb, f)

        return {"status": "success"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# downloads file from fs to client
from fastapi import Form

@app.post("/download-file")
async def download_file(api_key: str = Depends(verify_api_key), file_name: str = Form(...)):
    file_path = f"{workspace_path}/{file_name}"
    if os.path.isfile(file_path):
        return FileResponse(file_path, filename=file_name)
    else:
        raise HTTPException(status_code=404, detail="File not found")


# uploads file to supabase and returns SignedURL  
@app.post("/download-url")
async def upload_file(input: ULFileData, api_key: str = Depends(verify_api_key)):
    folder_path = f"{workspace_path}/{input.folder_name}"
    if not os.path.isdir(folder_path):
        raise HTTPException(status_code=400, detail="Folder does not exist")

    path_on_supabase = f"{input.folder_name}/{input.file_name}"

    file_path = f"{folder_path}/{input.file_name}"
    if not os.path.isfile(file_path):
        raise HTTPException(status_code=400, detail="File does not exist")

    try:

        # Delete the existing file if it exists
        supabase.storage.from_(input.bucket_name).remove([path_on_supabase])

        with open(file_path, 'rb') as f:
            supabase.storage.from_(input.bucket_name).upload(file=f, path=path_on_supabase)

        # Get the file URL
        #file_url = supabase.storage.from_(input.bucket_name).get_public_url(path_on_supabase)
        res = supabase.storage.from_(input.bucket_name).create_signed_url(path_on_supabase, input.expiry_duration)

        #return {"status": "success", "message": f"File {input.file_name} has been uploaded to {input.bucket_name}", "file_url": file_url, "signed_url": res}
        return {"status": "success", "message": f"File {input.file_name} has been uploaded to {input.bucket_name}", "signed_url": res}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

from moviepy.editor import VideoFileClip



@app.post("/upload-video-extract-audio")
async def upload_video_extract_audio(file: UploadFile = File(...), folder_name: str = Form(...), bucket_name: str = Form(...), api_key: str = Depends(verify_api_key)):
    folder_path = f"{workspace_path}/{folder_name}"
    if not os.path.isdir(folder_path):
        os.makedirs(folder_path, exist_ok=True)

    s3_client = boto3.client('s3')
    
    file_path = f"{folder_path}/{file.filename}"
    audio_filename = f"{file.filename.rsplit('.', 1)[0]}.mp3"
    audio_path = f"{folder_path}/{audio_filename}"

    try:
        # Save the uploaded video file
        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)

        # Extract audio from video
        video = VideoFileClip(file_path)
        video.audio.write_audiofile(audio_path)
        video.close()

        # Delete the original video file
        os.remove(file_path)

        # Upload the audio file to S3
        s3_client.upload_file(audio_path, bucket_name, f"{folder_name}/{audio_filename}")

        # Delete the audio file after upload
        os.remove(audio_path)

        return {"status": "success", "message": "Audio extracted and uploaded to S3", "filename": audio_filename, "folder": folder_name}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post('/v1/tokens')
async def count_tokens(request: Request, api_key: str = Depends(verify_api_key)):
    data = await request.json()
    if 'key' not in data:
        raise HTTPException(status_code=404, detail="Key not found")
    elif data['key'] != os.getenv('API_KEY'):
        raise HTTPException(status_code=401, detail="Invalid Key")
    text = data['messages']
    tokenizer = get_tokenizer('daryl149/llama-2-7b-chat-hf')
    # Tokenize the string
    tokens = tokenizer.tokenize(text)
    return {"token_count": len(tokens)}




MAX_CACHE_SIZE = 5  # Adjust depending on memory constraints
model_cache = OrderedDict()  # To maintain the order of insertion for eviction

def create_embeddings(input_texts, model_name):
    # Check if model is in cache
    if model_name in model_cache:
        model = model_cache[model_name]
        # Move the model to the end to show it's recently used
        model_cache.move_to_end(model_name)
    else:
        # If not in cache, load and add to cache
        #model = SentenceTransformer(model_name, device="cuda")
        model = SentenceTransformer("BAAI/bge-small-en-v1.5")
        if len(model_cache) >= MAX_CACHE_SIZE:
            # Evict the least recently used model (first key in OrderedDict)
            model_cache.popitem(last=False)
        model_cache[model_name] = model

    embeddings = model.encode(input_texts, normalize_embeddings=True)
    embeddings = embeddings.tolist()
    return embeddings

models_cache = OrderedDict()


def parse_youtube_transcript(url):
    loader = YoutubeLoader.from_youtube_url(url, add_video_info=True)
    res = loader.load()
    return res[0]


@app.post("/v1/split-text")
async def process_text(request_data: InputModel, api_key: str = Depends(verify_api_key)):
    chunk_size = request_data.chunks.chunk_size
    chunk_overlap = request_data.chunks.chunk_overlap
    return_full_text = request_data.return_full_text
    model = request_data.model
    input_type = request_data.type

    if input_type == "file":
        # Handle file URL
        file_url = request_data.url
        # Fetch the file
        response = requests.get(file_url)
        if response.status_code != 200:
            raise HTTPException(status_code=404, detail="File not found")

        # Get file type from the URL
        content_type = response.headers.get("Content-Type")
        MIME_MAP = {
            "application/pdf": "pdf",
            "application/msword": "doc",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
            "text/plain": "txt",
            "application/rtf": "rtf",
        }
        file_type = MIME_MAP.get(content_type, "unknown")

        # Parse the file based on its type
        text = parse_file(file_type, response)

    elif input_type == "youtube_transcript":
        youtube_text = parse_youtube_transcript(request_data.url)
        youtube_metadata = youtube_text.metadata
        text = youtube_text.page_content

    elif input_type == "text":
        # Handle plain text
        text = request_data.text

    else:
        raise HTTPException(status_code=400, detail="Unsupported input type")
    
    # Filter out Unicode escape characters
    text = re.sub(r'\\u[0-9A-Fa-f]{1,4}', '', text)

    # Remove null characters
    text = text.replace('\u0000', '')

    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        length_function=len,
        add_start_index=True,
    )
    documents = text_splitter.create_documents([text])

    # Combine all page_content into a single string
    whole_text = " ".join([doc.page_content for doc in documents])

    # Create chunks list
    if model is not None:
        chunks = [
            {
                "chunk": doc.page_content,
                "vector": create_embeddings(doc.page_content, model),
                "metadata": doc.metadata,
            }
            for doc in documents
        ]
    else:
        chunks = [
            {"chunk": doc.page_content, "metadata": doc.metadata}
            for doc in documents
        ]
    # Final output structure
    output = {
        "metadata": {
            "document_name": "input_text",
            "mime_type": "text/plain",
        },
        "chunks": chunks,
    }

    if return_full_text:
        output["text"] = whole_text

    if model:
        output["model"] = model

    return JSONResponse(content=output, status_code=200)

    
def parse_file(file_type, response):
    # Parse file based on the file type
    if file_type == "pdf":
        pdf_file = BytesIO(response.content)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        text = "\n".join(
            [pdf_reader.pages[i].extract_text() for i in range(len(pdf_reader.pages))]
        )

    elif file_type == "docx":
        doc_file = BytesIO(response.content)
        doc = Document(doc_file)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])

    elif file_type == "rtf":
        text = rtf_to_text(response.content.decode())


    elif file_type == "txt":
        text = response.text

    else:
        raise HTTPException(status_code=400, detail="Unsupported file type")

    return text


@app.post("/v1/split-pdf-paged")
async def process_pdf(request_data: InputModel):
    try:
        # Extract data from the request
        file_url = request_data.url
        chunk_size = request_data.chunks.chunk_size
        chunk_overlap = request_data.chunks.chunk_overlap
        model = request_data.model

        # Fetch the PDF file from the URL
        response = requests.get(file_url)
        if response.status_code != 200:
            raise HTTPException(status_code=404, detail="File not found")

        # Read the PDF file
        pdf_file = BytesIO(response.content)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        
        documents = []
        cumulative_length = 0
        for page_number, page in enumerate(pdf_reader.pages):
            text = page.extract_text()
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=chunk_size,
                chunk_overlap=chunk_overlap,
                length_function=len,
                add_start_index=True,
            )
            page_documents = text_splitter.create_documents([text])
            for doc in page_documents:
                doc.metadata['page_number'] = page_number + 1
                doc.metadata['start_index'] += cumulative_length
            cumulative_length += len(text)
            documents.extend(page_documents)

        # Create chunks list
        if model is not None:
            chunks = [
                {
                    "chunk": doc.page_content,
                    "start_index": doc.metadata['start_index'],
                    "page_number": doc.metadata['page_number'],
                    "vector": create_embeddings(doc.page_content, model),
                }
                for doc in documents
            ]
        else:
            chunks = [
                {
                    "chunk": doc.page_content,
                    "start_index": doc.metadata['start_index'],
                    "page_number": doc.metadata['page_number'],
                }
                for doc in documents
            ]

        return JSONResponse(content={"chunks": chunks}, status_code=200)
    except Exception as e:
        logging.error(f"Error processing PDF: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post('/v1/embeddings')
async def get_embeddings(request: Request, api_key: str = Depends(verify_api_key)):
    data = await request.json()
    model_name = data['model']
    input_texts = data['input']

    # Check if model is in cache
    if model_name in models_cache:
        model = models_cache[model_name]
        # Move the model to the end to show it's recently used
        models_cache.move_to_end(model_name)
    else:
        # If not in cache, load and add to cache
        #model = SentenceTransformer(model_name, device='cuda')
        model = SentenceTransformer("BAAI/bge-small-en-v1.5")
        if len(models_cache) >= MAX_CACHE_SIZE:
            # Evict the least recently used model (first key in OrderedDict)
            models_cache.popitem(last=False)
        models_cache[model_name] = model

    embeddings = model.encode(input_texts, normalize_embeddings=True)
    embeddings = embeddings.tolist()
    return {
        "data": [
            {
                "embedding": embeddings
            }
        ],
        "model": model_name
    }

@app.post("/transcribe")
async def transcribe(input: TranscribeInput, api_key: str = Depends(verify_api_key)):
    try:
        current_date = datetime.now().strftime("%d%m%Y")
        short_id = str(uuid.uuid4())[:8]
        job_name = short_id + "_" + current_date
        input_path = input.input_path
        output_path = input.output_path

        # Extract file extension
        file_type = input_path.split('.')[-1]
        
        s3_location = f"s3://townhallmeeting/{input_path}"
        #s3_location = f"s3://transcripts/{input_path}"

        # Get the base filename from the input path
        base_filename = os.path.basename(input_path)
        
        response = start_transcription(job_name, s3_location, file_type, output_path)
        
        if "TranscriptionJobStatus" in response and response["TranscriptionJobStatus"] == "IN_PROGRESS":
            return {"status": "Processing", "job_name": job_name, "json_file": f"{base_filename}.json", "srt_file": f"{base_filename}.srt"}
        else:
            return {"status": "error", "message": response}
    except Exception as e:
        logging.error(f"Unexpected error: {str(e)}")
        return {"status": "error", "message": str(e)}
    
def start_transcription(job_name, media_uri, type_file, output_path):
    base_filename = os.path.basename(media_uri)
    #filename = job_name + base_filename

    filetype = type_file

    try: 
        client = boto3.client('transcribe', region_name='us-gov-east-1')  # Specify your region here
        #client = boto3.client('transcribe', region_name='us-west-1') 

        response = client.start_transcription_job(
            TranscriptionJobName=str(job_name),
            LanguageCode='en-US',
            MediaFormat=filetype,
            Media={
                'MediaFileUri': media_uri,
            },
            OutputBucketName='townhallmeeting',
            #OutputBucketName='transcript',
            #OutputKey=f'{output_path}/{filename}.json',
            OutputKey=f'{output_path}/{base_filename}.json',
            Subtitles={
                'Formats': [
                    'srt',
                ],
                'OutputStartIndex': 1
            }
        )
        logging.info(f'{base_filename} Processing')
        return {
            "TranscriptionJobName": response['TranscriptionJob']['TranscriptionJobName'],
            "TranscriptionJobStatus": response['TranscriptionJob']['TranscriptionJobStatus']
        }
    except Exception as e:
        logging.error(f"Error starting transcription: {e}")
        return str(e)

class ResultInput(BaseModel):
    file_name: str
    status: str

@app.post("/transcriberesults")
async def receive_results(input: ResultInput, api_key: str = Depends(verify_api_key)):
    logging.info(f"Request payload: {input}")
    if input.status not in ["COMPLETED", "FAILED"]:
        raise HTTPException(status_code=400, detail="Invalid status. Must be 'Completed' or 'Failed'.")

    # Process the result here (e.g., log it, update a database, etc.)
    logging.info(f"Received result for file {input.file_name}: {input.status}")
    logging.info(input)

    return {"status": "success", "message": f"Result for file {input.file_name} received with status {input.status}"}

client = boto3.client('transcribe', region_name='us-gov-east-1')
#client = boto3.client('transcribe', region_name='us-west-1')

@app.get("/transcription-status")
async def transcription_status(job_name: str = Query(...)):
    try:
        print("AWS_ACCESS_KEY_ID:", os.getenv("AWS_ACCESS_KEY_ID"))
        print("AWS_SECRET_ACCESS_KEY:", os.getenv("AWS_SECRET_ACCESS_KEY"))
        print("AWS_SESSION_TOKEN:", os.getenv("AWS_SESSION_TOKEN"))
        response = client.get_transcription_job(
            TranscriptionJobName=job_name
        )
        return {"job name": job_name, "job_status": response['TranscriptionJob']['TranscriptionJobStatus']}
    except client.exceptions.BadRequestException as e:
        raise HTTPException(status_code=400, detail=str(e))
    except client.exceptions.LimitExceededException as e:
        raise HTTPException(status_code=429, detail=str(e))
    except client.exceptions.InternalFailureException as e:
        raise HTTPException(status_code=500, detail=str(e))
    except client.exceptions.ConflictException as e:
        raise HTTPException(status_code=409, detail=str(e))
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/environment")
async def get_file(token: str = Depends(verify_token)):
    return FileResponse('/home/ubuntu/automatenb/environment.txt')

@app.get("/api-health")
async def api_health():
    return {"status": "API is healthy"}



if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=5000)
